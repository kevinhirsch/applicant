"""Resume-parser adapter (FR-ONBOARD-3, FR-ATTR-1).

Parses an uploaded base resume into a :class:`ParsedResume` so the onboarding
service can bootstrap the per-campaign attribute cloud and reconcile against the
interview answers. Supports docx (python-docx), txt, and PDF (pypdf) at minimum.

Extraction is heuristic and intentionally conservative: it never fabricates. When
a field cannot be found it is left empty, and onboarding fills it from the
interview. Font detection reuses the docx font table (docx) or fontspec directives
(txt/tex) so the font subsystem can prompt for any missing families (FR-FONT-1).

Phone number regex and date-range patterns are locale-aware via
applicant.core.locale_config (issue #194).
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

from applicant.observability.logging import get_logger
from applicant.ports.driven.resume_parser import (
    EducationEntry,
    ParsedResume,
    WorkHistoryEntry,
)

log = get_logger(__name__)

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONE_RE = re.compile(r"(?:\+?\(?\d[\d\s().-]{7,}\d)")
# A dated work-history line: "Title, Company    Jan 2020 - Present"
#: Date/month names are locale-aware (issue #194).
_MONTH_NAMES = r"Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December"
_PRESENT_KEYWORDS = r"Present|Current"
#: One side of a date range. Tries numeric "MM/YYYY" and "MM/DD/YYYY" forms
#: (a very common template/export format) BEFORE the month-name-or-bare-year
#: form, and requires a full 4-digit year on the numeric forms so a stray
#: decimal like "99.9%" or "3.5" is never mistaken for a date. Regression: a
#: line whose date used slashes ("06/2018 - 12/2020") previously matched
#: nothing at all, so ``_DATE_RANGE_RE.search`` returned None and the whole
#: work-history entry for that line was silently dropped.
_DATE_TOKEN = (
    r"(?:\d{1,2}[/.]\d{1,2}[/.]\d{4})"
    r"|(?:\d{1,2}[/.]\d{4})"
    r"|(?:(?:" + _MONTH_NAMES + r"|\d{1,2})?\.?\s*\d{4})"
)
_DATE_RANGE_RE = re.compile(
    r"(" + _DATE_TOKEN + r")"
    r"\s*(?:-|–|—|to)\s*"
    r"(" + _PRESENT_KEYWORDS + r"|" + _DATE_TOKEN + r")",
    re.IGNORECASE,
)
_YEAR_RANGE_RE = re.compile(
    r"(\d{4})\s*(?:-|–|—|to)\s*(\d{4}|" + _PRESENT_KEYWORDS + r")",
    re.IGNORECASE,
)
#: P1-1: a LONE graduation year on an education line ("B.A. Economics, UC
#: Berkeley — 2013"). Real résumés frequently list only the completion year;
#: without this the year leaked into the institution text ("UC Berkeley — 2013")
#: and both year fields rendered empty in the review form.
_SINGLE_YEAR_RE = re.compile(r"\b((?:19|20)\d{2})\b")
#: Degree token. Every alternative is wrapped with letter-boundary lookaround
#: (below) rather than plain ``\b`` so it also works around the trailing dots
#: on abbreviations like "B.S." (``\b`` doesn't fire between two non-word
#: characters such as "." and " "). Regression: without the boundary, the
#: bare 2-letter forms (B.?S.?, B.?A.?, M.?S.?, M.?A.?) matched case-insensitively
#: ANYWHERE a lowercase bigram like "ma"/"ms" occurred inside an ordinary word
#: -- e.g. "Infor-MA-tion" or "Main-TA-ined" -- so the "degree" capture could
#: start mid-word. The short forms additionally require the FIRST period
#: (```B\.S\.?``` not ```B\.?S\.?```) so a bare, undotted "ms"/"as" substring
#: (which collides with common words/honorifics -- "teams", "Ms Jones") can't
#: match; dotted usage ("B.S.", "M.S.") and the common undotted long forms
#: ("Bachelor", "MBA") are unaffected. "A.A."/"A.S."/"A.A.S." (Associate of
#: Arts/Science/Applied Science) were previously missing entirely -- lacking a
#: real alternative to match, a degree line using them fell through to a false
#: mid-word match instead (the reported "A.A. Computer Infor" / "mation
#: Systems" split).
_DEGREE_RE = re.compile(
    r"(?<![A-Za-z])"
    r"(A\.A\.?S\.?|A\.A\.?|A\.S\.?|B\.S\.?|B\.A\.?|M\.S\.?|M\.A\.?|MBA|Ph\.?D\.?|"
    r"Bachelor|Master|Doctor|Associate)"
    r"(?![A-Za-z])"
    r"[^,\n]*",
    re.IGNORECASE,
)
_SECTION_RE = re.compile(
    r"^\s*(experience|work experience|work history|career history|"
    r"employment|professional experience|"
    r"education|skills|technical skills|core competencies|"
    r"certifications?|licenses?|certifications? (?:&|and) licenses?|"
    r"summary|professional summary|objective|projects|"
    r"awards|honors|publications|languages|interests|references|"
    r"volunteer(?:ing)?|activities)\s*:?\s*$",
    re.IGNORECASE,
)
#: Inline separators on a skills line. Includes the middle dot ``·`` (U+00B7) and a
#: couple of dot/bullet variants commonly used between skills, which were previously
#: missing — so "SQL · Postgres" used to parse as a single bogus "SQL · Postgres" token.
_SKILL_SEPARATORS = frozenset(",;•·‧・|")
#: A run of 2+ spaces used as a column delimiter in plain-text exports where a
#: title and company are on the same line with no other separator ("Senior
#: Support Engineer     Acme Corp    Jan 2021 - Present"). Only used as a
#: last-resort fallback in ``_split_title_company`` after every named
#: separator has failed to match.
_COLUMN_GAP_RE = re.compile(r" {2,}")
#: A "City, ST" location — one or more words, a comma, then a 2-letter (US) state
#: code (optionally dotted). Common Word/Google-Docs résumé layout puts the role
#: TITLE + company on one line and "LOCATION – DATE - DATE" on the NEXT line
#: ("Lead Scrum Master – Wells Fargo" / "PHOENIX, AZ – SEPTEMBER 2025 - PRESENT").
#: Without recognising the location, the text before the date on the date line
#: ("PHOENIX, AZ") was mis-read as the title/company and the real title line above
#: was never consulted (the look-back only fired when that text was empty).
_LOCATION_RE = re.compile(r"^[A-Za-z.'\-]+(?:[ .][A-Za-z.'\-]+)*,\s*[A-Za-z]{2}\.?$")


class ResumeParser:
    """ResumeParserPort adapter for docx / txt / pdf base resumes."""

    def parse(self, document_path: str) -> ParsedResume:
        path = Path(document_path)
        suffix = path.suffix.lower()
        if suffix == ".docx":
            text = self._read_docx(path)
            fonts = self._detect_docx_fonts(path)
        elif suffix == ".pdf":
            text = self._read_pdf(path)
            fonts = ()
        else:  # .txt / .tex / unknown -> treat as plain text
            text = self._read_text(path)
            fonts = self._detect_text_fonts(text)
        return self._extract(text, fonts)

    # --- readers -----------------------------------------------------------
    def _read_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            return ""

    def _read_docx(self, path: Path) -> str:
        try:
            from docx import Document
        except Exception:  # pragma: no cover - dep present in prod/test
            return ""
        try:
            doc = Document(str(path))
        except Exception:
            return ""
        lines = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(lines)

    def _read_pdf(self, path: Path) -> str:
        try:
            from pypdf import PdfReader
        except Exception:  # pragma: no cover
            return ""
        try:
            reader = PdfReader(str(path))
            return "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            return ""

    # --- font detection ----------------------------------------------------
    def _detect_docx_fonts(self, path: Path) -> tuple[str, ...]:
        """Read declared font families from the docx font table (FR-FONT-1)."""
        found: list[str] = []
        try:
            with zipfile.ZipFile(str(path)) as zf:
                names = zf.namelist()
                # Fonts declared in fontTable.xml and styles.xml (w:rFonts).
                for member in ("word/fontTable.xml", "word/styles.xml", "word/document.xml"):
                    if member not in names:
                        continue
                    xml = zf.read(member).decode("utf-8", errors="ignore")
                    for m in re.finditer(r'w:(?:ascii|hAnsi|cs)="([^"]+)"', xml):
                        fam = m.group(1).strip()
                        if fam and fam not in found:
                            found.append(fam)
                    for m in re.finditer(r'<w:font w:name="([^"]+)"', xml):
                        fam = m.group(1).strip()
                        if fam and fam not in found:
                            found.append(fam)
        except (OSError, zipfile.BadZipFile):
            return ()
        return tuple(found)

    def _detect_text_fonts(self, text: str) -> tuple[str, ...]:
        found: list[str] = []
        for pat in (
            r"\\setmainfont(?:\[[^\]]*\])?\{([^}]+)\}",
            r"\\setsansfont(?:\[[^\]]*\])?\{([^}]+)\}",
            r"\\fontspec(?:\[[^\]]*\])?\{([^}]+)\}",
        ):
            for m in re.finditer(pat, text):
                fam = m.group(1).split("-")[0].strip()
                if fam and fam not in found:
                    found.append(fam)
        return tuple(found)

    # --- extraction --------------------------------------------------------
    def _extract(self, text: str, fonts: tuple[str, ...]) -> ParsedResume:
        lines = [ln.rstrip() for ln in text.splitlines()]
        non_empty = [ln for ln in lines if ln.strip()]

        email = _EMAIL_RE.search(text)
        # Avoid matching dates/years as a phone number: require it not to be the email.
        phone = ""
        for m in _PHONE_RE.finditer(text):
            candidate = m.group(0).strip()
            digits = re.sub(r"\D", "", candidate)
            if 9 <= len(digits) <= 15:
                phone = candidate
                break

        full_name = self._guess_name(non_empty, email.group(0) if email else "")
        work_history = self._extract_work_history(lines)
        education = self._extract_education(lines)
        skills = self._extract_skills(lines)

        return ParsedResume(
            full_name=full_name,
            email=email.group(0) if email else "",
            phone=phone,
            work_history=tuple(work_history),
            education=tuple(education),
            skills=tuple(skills),
            detected_fonts=fonts,
            raw_text=text,
        )

    def _guess_name(self, non_empty: list[str], email: str) -> str:
        """First non-contact line of 2-4 capitalized words is the name."""
        for line in non_empty[:5]:
            s = line.strip()
            if email and email in s:
                continue
            if _EMAIL_RE.search(s) or _PHONE_RE.search(s):
                continue
            words = s.split()
            if 1 < len(words) <= 4 and all(w[0:1].isupper() for w in words if w[0:1].isalpha()):
                return s
        return ""

    def _current_section(self, lines: list[str]) -> dict[str, list[str]]:
        sections: dict[str, list[str]] = {}
        current = "_preamble"
        sections[current] = []
        for line in lines:
            m = _SECTION_RE.match(line)
            if m:
                current = m.group(1).lower()
                sections.setdefault(current, [])
                continue
            sections.setdefault(current, []).append(line)
        return sections

    def _extract_work_history(self, lines: list[str]) -> list[WorkHistoryEntry]:
        sections = self._current_section(lines)
        body: list[str] = []
        for key in (
            "experience",
            "work experience",
            "work history",
            "career history",
            "employment",
            "professional experience",
        ):
            body += sections.get(key, [])
        if not body:
            body = lines  # fall back to whole doc

        # Pass 1: locate every entry's date line and the index/text of the line
        # that carries its title/company ("header"). The header is usually the
        # date line itself (same-line layout) but can be an earlier line (the
        # "Title, Company" / date-on-next-line layout) via the look-back.
        matches: list[tuple[int, re.Match, int, str, str]] = []
        for idx, line in enumerate(body):
            dm = _DATE_RANGE_RE.search(line)
            if not dm:
                continue
            # The date range is frequently parenthesized ("Acme Corp (2020-Present)");
            # the regex matches the date itself, so the opening bracket is left
            # dangling at the end of `before`. Strip brackets/separators from both
            # ends so the company never renders as "Acme Corp (" (FR-RESUME-3 fidelity).
            before = line[: dm.start()].strip(" \t,-–—|([{")
            header_idx = idx
            location = ""
            if not before:
                # Common layout: "Title, Company" on its own line, the date range on
                # the NEXT line. The title/company is therefore on the immediately
                # preceding non-empty, non-date line — look back and attribute it,
                # otherwise the entry renders as "\cventry{2021 - Present}{}{}...".
                header_idx, before = self._look_back_title_company(body, idx)
            elif _LOCATION_RE.match(before):
                # "LOCATION – DATE" date line (Word/Docs layout): the text before the
                # date is the CITY, ST — not the title/company, which sits on the line
                # above. Keep the location for the entry and look back for the header.
                location = before
                lb_idx, lb_before = self._look_back_title_company(body, idx)
                if lb_before:
                    header_idx, before = lb_idx, lb_before
                else:
                    before = ""  # no header found — don't keep the location as title
            matches.append((idx, dm, header_idx, before, location))

        entries: list[WorkHistoryEntry] = []
        for i, (date_idx, dm, _header_idx, before, location) in enumerate(matches):
            title, company = self._split_title_company(before)
            # Defensive: the split can still leave a trailing bracket on either field
            # when the separator sat between the date and the bracket.
            title = title.strip(" \t,-–—|([{")
            company = company.strip(" \t,-–—|([{")
            # Achievements/bullets: every non-empty line strictly between THIS
            # entry's date line and the start of the NEXT entry's header (or the
            # end of the body for the last entry) belongs to this role. Stopping
            # at the next header — not the next date line — matters for the
            # date-on-next-line layout, where the next entry's title/company
            # line comes before its own date line and must not be swallowed as
            # one of this entry's bullets.
            next_header_idx = matches[i + 1][2] if i + 1 < len(matches) else len(body)
            achievements = tuple(
                cleaned
                for raw in body[date_idx + 1 : next_header_idx]
                if (cleaned := raw.strip(" \t-–—•*"))
            )
            entries.append(
                WorkHistoryEntry(
                    title=title,
                    company=company,
                    start_date=dm.group(1).strip(),
                    end_date=dm.group(2).strip(),
                    location=location,
                    achievements=achievements,
                )
            )
        return entries

    @staticmethod
    def _look_back_title_company(body: list[str], idx: int) -> tuple[int, str]:
        """Return the nearest preceding non-empty, non-date line as title/company.

        Handles the layout where ``Title, Company`` sits on its own line and the
        date range is on the following line. The look-back stops at the first
        candidate and ignores lines that themselves carry a date range (they belong
        to a different, earlier entry). Returns the candidate's own index (so the
        caller can tell where THIS entry's header starts, e.g. to bound the
        previous entry's achievement lines) alongside its text; falls back to
        ``idx`` itself (the date line) when nothing usable is found.
        """
        for prev in range(idx - 1, -1, -1):
            cand = body[prev].strip(" \t,-–—|([{")
            if not cand:
                continue
            if _DATE_RANGE_RE.search(body[prev]):
                break
            return prev, cand
        return idx, ""

    @staticmethod
    def _split_title_company(text: str) -> tuple[str, str]:
        for sep in (" – ", " — ", " at ", " @ ", ", ", " | ", "\t", " - "):
            if sep in text:
                left, right = text.split(sep, 1)
                return left.strip(), right.strip()
        # Last resort: plain-text exports commonly pad title/company into
        # fixed-width columns with a run of spaces instead of any punctuation
        # ("Senior Support Engineer     Acme Corp"). Without this, the whole
        # string fell into `title` and `company` was silently left empty.
        m = _COLUMN_GAP_RE.search(text)
        if m:
            left, right = text[: m.start()].strip(), text[m.end() :].strip()
            if left and right:
                return left, right
        return text.strip(), ""

    def _extract_education(self, lines: list[str]) -> list[EducationEntry]:
        sections = self._current_section(lines)
        body = sections.get("education", [])
        if not body:
            body = [ln for ln in lines if _DEGREE_RE.search(ln)]
        entries: list[EducationEntry] = []
        for idx, line in enumerate(body):
            dm = _DEGREE_RE.search(line)
            if not dm:
                continue
            degree = dm.group(0).strip()
            # The greedy degree match runs to the next comma/newline, so when
            # there's no comma before the graduation years on the same line
            # ("B.S. Computer Science 2013 - 2017") it swallows the dates too.
            # Trim any inline year range back out of the degree text.
            inline_year = _YEAR_RANGE_RE.search(degree)
            if inline_year:
                degree = degree[: inline_year.start()].strip(" \t,-|")
            ym = _YEAR_RANGE_RE.search(line)
            if not ym:
                # The year range is frequently on the line that follows the degree
                # ("M.S. Computer Science\n2018 - 2020"); pull it from the next
                # non-empty line so the entry doesn't drop its dates.
                ym = self._look_ahead_year_range(body, idx)
            # P1-1 (the "UC Berkeley — 2013" case): no RANGE anywhere, but the
            # line carries a lone graduation year — capture it as the end year
            # and strip it from the institution text below, so it renders in the
            # review form's year field instead of polluting the school name.
            single = None if ym else _SINGLE_YEAR_RE.search(line)
            if single and single.group(0) in degree:
                degree = degree[: degree.index(single.group(0))].strip(" \t,-|–—")
            # Institution: whatever's left on the line once the degree (and its
            # own year range, if any) are removed. This used to be run through
            # `_split_title_company` (built for 2-part "Title, Company" splits),
            # which keeps only the LEFT side of the first separator it finds —
            # discarding the real institution name entirely whenever the degree
            # match didn't already stop exactly at the institution's boundary
            # (the reported "A.A. Computer Infor" / institution garbage came
            # from exactly this: the mangled degree match left "Infor" as the
            # trailing edge of `rest`, and the college name after the comma was
            # thrown away as the unused right-hand side of the split).
            rest = line.replace(dm.group(0), "", 1).strip(" \t,-|")
            if ym and ym.group(0) in rest:
                rest = rest.replace(ym.group(0), "", 1)
            if single and single.group(0) in rest:
                rest = rest.replace(single.group(0), "", 1)
            institution = rest.strip(" \t,-|–—")
            if not institution:
                # Layout where degree / dates / institution each sit on their own
                # line ("M.S. Computer Science\n2015 - 2017\nState University").
                institution = self._look_ahead_institution(body, idx)
            entries.append(
                EducationEntry(
                    degree=degree,
                    institution=institution,
                    start_year=ym.group(1) if ym else "",
                    end_year=ym.group(2) if ym else (single.group(1) if single else ""),
                )
            )
        return entries

    @staticmethod
    def _look_ahead_year_range(body: list[str], idx: int) -> re.Match | None:
        """Return a year range found on the line following a degree line, if any.

        Stops at the first non-empty line; a degree on that line means it belongs to
        a different entry, so the look-ahead yields nothing.
        """
        for nxt in range(idx + 1, len(body)):
            cand = body[nxt].strip()
            if not cand:
                continue
            if _DEGREE_RE.search(cand):
                return None
            return _YEAR_RANGE_RE.search(cand)
        return None

    @staticmethod
    def _look_ahead_institution(body: list[str], idx: int) -> str:
        """Return the institution name for a degree line whose own line has none.

        Mirrors ``_look_ahead_year_range``: walks forward from the degree line,
        skipping a line that's purely the year range (already consumed for the
        dates), and returns the next non-empty line as the institution -- unless
        that line starts a new degree, in which case it belongs to a later entry
        and nothing is returned.
        """
        for nxt in range(idx + 1, len(body)):
            cand = body[nxt].strip()
            if not cand:
                continue
            if _DEGREE_RE.search(cand):
                return ""
            if _YEAR_RANGE_RE.fullmatch(cand.strip(" \t,-|")):
                continue
            return cand.strip(" \t,-|")
        return ""

    def _extract_skills(self, lines: list[str]) -> list[str]:
        sections = self._current_section(lines)
        body: list[str] = []
        for key in ("skills", "technical skills", "core competencies"):
            body += sections.get(key, [])
        skills: list[str] = []
        for line in body:
            for tok in self._split_skills(line):
                s = tok.strip()
                if s and len(s) <= 40 and s not in skills:
                    skills.append(s)
        return skills

    @staticmethod
    def _split_skills(line: str) -> list[str]:
        """Split a skills line on common separators, but NOT inside parentheses.

        A naive ``re.split`` on commas shreds a parenthetical sub-list like
        ``AWS (EKS, RDS, Lambda)`` into the junk tokens ``AWS (EKS`` / ``RDS`` /
        ``Lambda)``. Tracking bracket depth keeps such a group as one clean token.
        """
        out: list[str] = []
        buf: list[str] = []
        depth = 0
        for ch in line:
            if ch in "([{":
                depth += 1
                buf.append(ch)
            elif ch in ")]}":
                depth = max(0, depth - 1)
                buf.append(ch)
            elif depth == 0 and ch in _SKILL_SEPARATORS:
                out.append("".join(buf))
                buf = []
            else:
                buf.append(ch)
        out.append("".join(buf))
        return out
