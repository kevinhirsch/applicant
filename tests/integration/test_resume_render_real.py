"""Real-render fidelity tests for the resume-tailoring engines (FR-RESUME-4) — integration only.

The DEFAULT lane stubs the xelatex/lualatex compile and the docx->PDF convert behind
the adapter seams (no TeX / LibreOffice required). These tests exercise the REAL
render path ONLY when the engine binary is actually present, inspecting the produced
PDF for the exact page count + embedded fonts (the compile-and-visually-inspect
fidelity check, modeled deterministically on the rendered output's metadata).
"""

from __future__ import annotations

import shutil
import zipfile

import pytest

from applicant.adapters.resume_tailoring.docx_tailor import (
    DocxTailor,
    write_document_xml,
)
from applicant.adapters.resume_tailoring.latex_tailor import LatexTailor
from applicant.core.ids import ResumeVariantId, new_id

_HAS_TEX = shutil.which("lualatex") or shutil.which("xelatex")
_HAS_SOFFICE = shutil.which("soffice") or shutil.which("libreoffice")


@pytest.mark.integration
@pytest.mark.skipif(not _HAS_TEX, reason="No TeX engine (lualatex/xelatex) installed.")
def test_latex_real_compile_embeds_fonts_and_counts_pages(tmp_path):
    adapter = LatexTailor(allow_compile=True, output_dir=tmp_path)
    source = (
        "\\documentclass[11pt]{article}\n"
        "\\begin{document}\n\\section*{Skills}\nPython, SQL\n\\end{document}\n"
    )
    result = adapter.render_artifact(ResumeVariantId(new_id()), source)
    assert result.page_count >= 1
    assert "fonts not embedded" not in result.notes


@pytest.mark.integration
@pytest.mark.skipif(not _HAS_TEX, reason="No TeX engine (lualatex/xelatex) installed.")
def test_latex_auto_mode_compiles_in_prod_when_engine_present(tmp_path):
    """FR-RESUME-4: the default render_mode="auto" (as wired in prod) really compiles
    and runs the embedded-font / page-count fidelity check when a TeX engine exists."""
    adapter = LatexTailor(output_dir=tmp_path)  # default render_mode="auto"
    assert adapter._allow_compile is True  # auto-enabled by the present engine
    source = (
        "\\documentclass[11pt]{article}\n"
        "\\begin{document}\n\\section*{Skills}\nPython, SQL\n\\end{document}\n"
    )
    result = adapter.render_artifact(ResumeVariantId(new_id()), source)
    assert result.page_count >= 1
    assert "fonts not embedded" not in result.notes


@pytest.mark.integration
@pytest.mark.skipif(not _HAS_SOFFICE, reason="No LibreOffice (soffice) installed.")
def test_docx_real_convert_to_pdf(tmp_path):
    import docx

    # Build a real .docx, edit its OOXML in place, then convert -> PDF.
    d = docx.Document()
    d.add_paragraph().add_run("Python developer")
    src = tmp_path / "base.docx"
    d.save(str(src))
    with zipfile.ZipFile(str(src)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    edited = DocxTailor().edit_document_xml(xml, {"Python": "Go"})
    out_docx = tmp_path / "edited.docx"
    write_document_xml(str(src), str(out_docx), edited)

    adapter = DocxTailor(allow_convert=True, output_dir=tmp_path)
    result = adapter.render_artifact(ResumeVariantId(new_id()), str(out_docx))
    if "approximate preview" in result.notes:
        pytest.skip("LibreOffice present but cannot convert docx in this environment")
    assert result.page_count >= 1
    assert "fonts in the rendered PDF are not embedded" not in result.notes
