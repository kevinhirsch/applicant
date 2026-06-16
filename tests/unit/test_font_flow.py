"""Font upload/management flow tests (FR-FONT-1/2).

Covers required-font detection, missing-font reporting, install + runtime
cache-refresh (stubbed fc-cache boundary; confined real FS copy), and
confirm-once-installed. No fontconfig / system-wide writes.
"""

from __future__ import annotations

import pytest

from applicant.adapters.fonts import font_installer as fi_mod
from applicant.adapters.fonts.font_installer import FontInstaller
from applicant.application.services.font_service import FontService


@pytest.fixture
def installer(tmp_path) -> FontInstaller:
    return FontInstaller(install_root=str(tmp_path / "fonts"))


@pytest.fixture
def service(installer) -> FontService:
    return FontService(installer)


def test_detect_required_fonts(installer, tmp_path):
    doc = tmp_path / "resume.tex"
    doc.write_text(
        "\\setmainfont{Lato-Lig}\n\\fontspec{Inconsolata}\n", encoding="utf-8"
    )
    found = installer.detect_required_fonts(str(doc))
    assert "Lato" in found
    assert "Inconsolata" in found


def _write_docx_with_font(path, family: str) -> None:
    """Build a tiny real .docx (a zip) declaring ``family`` via python-docx."""
    import docx

    d = docx.Document()
    run = d.add_paragraph().add_run("Jane Smith")
    run.font.name = family
    d.save(str(path))


def test_detect_required_fonts_reads_real_docx_font_table(installer, tmp_path):
    """FR-FONT-1: a real .docx (zip) declares fonts in word/*.xml, not as UTF-8 text.

    Before the fix ``detect_required_fonts`` read the zip bytes as UTF-8 and scanned
    for LaTeX directives only, so it found nothing for a genuine docx.
    """
    doc = tmp_path / "resume.docx"
    _write_docx_with_font(doc, "Segoe UI")
    found = installer.detect_required_fonts(str(doc))
    assert "Segoe UI" in found


def test_missing_fonts_reports_docx_declared_font(service, tmp_path):
    """FR-FONT-1: a docx-declared, non-system font surfaces in the missing list."""
    doc = tmp_path / "resume.docx"
    _write_docx_with_font(doc, "Segoe UI")
    report = service.report_for_document(str(doc))
    assert "Segoe UI" in report.missing


def test_missing_font_reporting_excludes_bundled_and_system(service, tmp_path):
    doc = tmp_path / "resume.tex"
    doc.write_text(
        "\\setmainfont{Lato}\n\\setsansfont{Calibri}\n\\fontspec{Inconsolata}\n",
        encoding="utf-8",
    )
    report = service.report_for_document(str(doc))
    assert "Inconsolata" in report.missing
    assert "Lato" not in report.missing  # bundled
    assert "Calibri" not in report.missing  # system font


def test_install_copies_into_confined_dir_and_refreshes_cache(installer, tmp_path):
    font = tmp_path / "Inconsolata.ttf"
    font.write_bytes(b"\x00\x01font-bytes")
    assert installer.cache_refresh_count == 0
    status = installer.install_font(str(font), "Inconsolata")
    assert status.installed is True
    # Real, confined copy happened (no system-wide write).
    installed_dir = tmp_path / "fonts"
    assert any(p.name.startswith("Inconsolata") for p in installed_dir.iterdir())
    # The (stubbed) fc-cache refresh ran exactly once.
    assert installer.cache_refresh_count == 1


def test_install_cannot_escape_confined_dir(installer, tmp_path):
    font = tmp_path / "evil.ttf"
    font.write_bytes(b"\x00")
    # A path-traversal name is sanitized; the file stays inside the confined dir.
    installer.install_font(str(font), "../../../etc/evil")
    installed_dir = tmp_path / "fonts"
    files = list(installed_dir.iterdir())
    assert files and all(p.parent == installed_dir for p in files)


def test_confirm_once_installed(service, tmp_path):
    font = tmp_path / "Inconsolata.ttf"
    font.write_bytes(b"\x00\x01")
    report = service.install(str(font), "Inconsolata")
    assert "Inconsolata" in report.installed
    # No longer reported missing on a fresh report.
    again = service.report_for_fonts(["Inconsolata"])
    assert again.missing == []


def test_fc_cache_runs_real_against_confined_dir_when_present(installer, tmp_path, monkeypatch):
    """FR-FONT-2: real ``fc-cache -f <install_root>`` shells out when present.

    Mocks the subprocess + PATH lookup (the marked seam) so the default lane needs
    NO fontconfig, but asserts the real call would target the confined dir.
    """
    calls: list[list[str]] = []

    def fake_which(name: str) -> str | None:
        return "/usr/bin/fc-cache" if name == "fc-cache" else None

    def fake_run(cmd, **kwargs):
        calls.append(cmd)

        class _R:
            returncode = 0

        return _R()

    monkeypatch.setattr(fi_mod.shutil, "which", fake_which)
    monkeypatch.setattr(fi_mod.subprocess, "run", fake_run)

    font = tmp_path / "Inconsolata.ttf"
    font.write_bytes(b"\x00\x01")
    installer.install_font(str(font), "Inconsolata")

    assert installer.fc_cache_invoked is True
    assert len(calls) == 1
    # fc-cache -f <confined install_root>, never system-wide.
    cmd = calls[0]
    assert cmd[0] == "/usr/bin/fc-cache"
    assert cmd[1] == "-f"
    confined = str((tmp_path / "fonts").resolve())
    assert cmd[2] == confined
    assert installer.last_cache_dir == confined


def test_fc_cache_graceful_noop_when_absent(installer, tmp_path, monkeypatch):
    """FR-FONT-2: no fontconfig -> graceful no-op (counted), no subprocess."""

    def fake_which(name: str) -> str | None:
        return None  # fc-cache not on PATH

    def boom(*a, **k):  # subprocess must NOT be called when fc-cache is absent
        raise AssertionError("fc-cache should not be invoked when absent")

    monkeypatch.setattr(fi_mod.shutil, "which", fake_which)
    monkeypatch.setattr(fi_mod.subprocess, "run", boom)

    font = tmp_path / "Inconsolata.ttf"
    font.write_bytes(b"\x00\x01")
    status = installer.install_font(str(font), "Inconsolata")

    assert status.installed is True  # install still succeeds
    assert installer.cache_refresh_count == 1  # refresh attempted (no-op)
    assert installer.fc_cache_invoked is False  # but no real fc-cache ran
    assert installer.last_cache_dir == str((tmp_path / "fonts").resolve())


def test_installs_persist_across_restart(tmp_path):
    root = str(tmp_path / "fonts")
    inst1 = FontInstaller(install_root=root)
    font = tmp_path / "Inconsolata.ttf"
    font.write_bytes(b"\x00\x01")
    inst1.install_font(str(font), "Inconsolata")
    # New installer over same confined dir = restart; rescans installed fonts.
    inst2 = FontInstaller(install_root=root)
    assert inst2.missing_fonts(["Inconsolata"]) == []
